import os
import sys
from importlib import import_module, reload
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from httpx import ASGITransport, AsyncClient


def build_docx_bytes(lines: list[str]) -> bytes:
    document = Document()
    for line in lines:
        document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def clear_provider_env(monkeypatch: pytest.MonkeyPatch) -> None:
    for key in [
        "OPENAI_API_KEY",
        "DEEPSEEK_API_KEY",
        "DOUBAO_API_KEY",
        "DOUBAO_MODELS",
        "TONGYI_API_KEY",
        "ZHIPU_API_KEY",
    ]:
        monkeypatch.delenv(key, raising=False)


def load_create_app():
    module_name = "app.main"
    if module_name in sys.modules:
        module = reload(sys.modules[module_name])
    else:
        module = import_module(module_name)
    return module.create_app


def write_env_file(tmp_path, content: str) -> None:
    (tmp_path / ".env").write_text(content, encoding="utf-8")


def test_create_app_fails_when_default_doubao_key_missing(tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    clear_provider_env(monkeypatch)
    monkeypatch.setenv("DATABASE_URL", f"sqlite+aiosqlite:///{tmp_path}/app.db")
    monkeypatch.setenv("STORAGE_ROOT", str(tmp_path / "storage"))
    monkeypatch.setenv("DEFAULT_PROVIDER", "doubao")
    create_app = load_create_app()

    with pytest.raises(ValueError, match="DOUBAO_API_KEY"):
        create_app()


def test_create_app_fails_when_default_doubao_models_missing(tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    clear_provider_env(monkeypatch)
    monkeypatch.setenv("DATABASE_URL", f"sqlite+aiosqlite:///{tmp_path}/app.db")
    monkeypatch.setenv("STORAGE_ROOT", str(tmp_path / "storage"))
    monkeypatch.setenv("DEFAULT_PROVIDER", "doubao")
    monkeypatch.setenv("DOUBAO_API_KEY", "test-key")
    create_app = load_create_app()

    with pytest.raises(ValueError, match="DOUBAO_MODELS"):
        create_app()


@pytest.mark.asyncio
async def test_models_endpoint_uses_configured_doubao_models_and_first_as_default(tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    clear_provider_env(monkeypatch)
    monkeypatch.setenv("DATABASE_URL", f"sqlite+aiosqlite:///{tmp_path}/app.db")
    monkeypatch.setenv("STORAGE_ROOT", str(tmp_path / "storage"))
    monkeypatch.setenv("DEFAULT_PROVIDER", "doubao")
    monkeypatch.setenv("DOUBAO_API_KEY", "test-key")
    monkeypatch.setenv(
        "DOUBAO_MODELS",
        "doubao-seed-1-6-flash-250715, doubao-seed-2-0-pro-250215, doubao-seed-1-6-flash-250715 ,, ",
    )

    create_app = load_create_app()
    app = create_app()

    async with AsyncClient(
        transport=ASGITransport(app=app),
        base_url="http://testserver",
    ) as client:
        response = await client.get("/api/models")

    assert response.status_code == 200
    payload = response.json()
    assert payload == [
        {
            "provider": "doubao",
            "name": "doubao-seed-1-6-flash-250715",
            "label": "doubao-seed-1-6-flash-250715",
            "target_langs": ["zh", "en", "th", "ar"],
            "default": True,
        },
        {
            "provider": "doubao",
            "name": "doubao-seed-2-0-pro-250215",
            "label": "doubao-seed-2-0-pro-250215",
            "target_langs": ["zh", "en", "th", "ar"],
            "default": False,
        }
    ]


@pytest.mark.asyncio
async def test_models_endpoint_merges_all_doubao_api_models_listed_in_env_comments(tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    clear_provider_env(monkeypatch)
    write_env_file(
        tmp_path,
        "\n".join(
            [
                "DEFAULT_PROVIDER=doubao",
                "DOUBAO_API_KEY=test-key",
                "DOUBAO_MODELS=doubao-seed-1-6-flash-250715",
                "#deepseek-v4-flash-260425",
                "# doubao-seed-2-0-pro-260215",
                "#doubao-seed-character-251128",
                "#glm-4-7-251222",
            ]
        ),
    )
    monkeypatch.setenv("DATABASE_URL", f"sqlite+aiosqlite:///{tmp_path}/app.db")
    monkeypatch.setenv("STORAGE_ROOT", str(tmp_path / "storage"))

    create_app = load_create_app()
    app = create_app()

    async with AsyncClient(
        transport=ASGITransport(app=app),
        base_url="http://testserver",
    ) as client:
        response = await client.get("/api/models")

    assert response.status_code == 200
    payload = response.json()
    assert [item["name"] for item in payload] == [
        "doubao-seed-1-6-flash-250715",
        "deepseek-v4-flash-260425",
        "doubao-seed-2-0-pro-260215",
        "doubao-seed-character-251128",
        "glm-4-7-251222",
    ]
    assert payload[0]["default"] is True
    assert all(item["provider"] == "doubao" for item in payload)


@pytest.mark.asyncio
async def test_upload_translate_and_download_round_trip_with_real_doubao(tmp_path, monkeypatch):
    doubao_api_key = os.environ.get("DOUBAO_API_KEY")
    if not doubao_api_key:
        pytest.skip("DOUBAO_API_KEY is required for real Doubao integration testing.")
    doubao_models = os.environ.get("DOUBAO_MODELS", "doubao-seed-1-6-flash-250715")
    doubao_model = doubao_models.split(",")[0].strip()

    monkeypatch.chdir(tmp_path)
    clear_provider_env(monkeypatch)
    monkeypatch.setenv("DATABASE_URL", f"sqlite+aiosqlite:///{tmp_path}/app.db")
    monkeypatch.setenv("STORAGE_ROOT", str(tmp_path / "storage"))
    monkeypatch.setenv("DEFAULT_PROVIDER", "doubao")
    monkeypatch.setenv("DOUBAO_API_KEY", doubao_api_key)
    monkeypatch.setenv("DOUBAO_MODELS", doubao_models)

    create_app = load_create_app()
    app = create_app()

    async with AsyncClient(
        transport=ASGITransport(app=app),
        base_url="http://testserver",
    ) as client:
        upload_lines = [
            "【第1集】",
            "艾米丽（局促，低声）：อีธาน พอแค่นี้ดีกว่าไหม ที่นี่คนเยอะเกินไป",
            "△ 场景描述",
        ]
        files = {
            "file": (
                "sample.docx",
                build_docx_bytes(upload_lines),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        data = {"title": "火花瞬间燃点"}

        upload_response = await client.post("/api/scripts", files=files, data=data)

        assert upload_response.status_code == 201
        payload = upload_response.json()
        assert payload["title"] == "火花瞬间燃点"
        assert payload["line_count"] == 3

        script_id = payload["script_id"]
        translate_response = await client.post(
            f"/api/scripts/{script_id}/translate",
            json={"target_lang": "zh", "provider": "doubao", "model": doubao_model},
        )

        assert translate_response.status_code == 202
        version_id = translate_response.json()["version_id"]

        translation_response = await client.get(f"/api/translations/{version_id}")
        assert translation_response.status_code == 200
        translation_payload = translation_response.json()
        assert translation_payload["status"] == "done"
        assert translation_payload["rendered_lines"][1].startswith("艾米丽（局促，低声）：")
        assert translation_payload["rendered_lines"][1].endswith(")")
        assert "(" in translation_payload["rendered_lines"][1]

        download_response = await client.get(f"/api/translations/{version_id}/download")
        assert download_response.status_code == 200
        assert download_response.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


@pytest.mark.asyncio
async def test_cleaned_script_upload_history_download_and_migration_preserve_existing_data(tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    clear_provider_env(monkeypatch)
    monkeypatch.setenv("DATABASE_URL", f"sqlite+aiosqlite:///{tmp_path}/app.db")
    monkeypatch.setenv("STORAGE_ROOT", str(tmp_path / "storage"))
    monkeypatch.setenv("DEFAULT_PROVIDER", "doubao")
    monkeypatch.setenv("DOUBAO_API_KEY", "test-key")
    monkeypatch.setenv("DOUBAO_MODELS", "doubao-seed-1-6-flash-250715")

    create_app = load_create_app()
    app = create_app()

    async with AsyncClient(
        transport=ASGITransport(app=app),
        base_url="http://testserver",
    ) as client:
        existing_files = {
            "file": (
                "original.docx",
                build_docx_bytes(["艾米丽：hello"]),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        existing_response = await client.post("/api/scripts", files=existing_files, data={"title": "既有剧本"})
        assert existing_response.status_code == 201

        files = {
            "file": (
                "translated.docx",
                build_docx_bytes(
                    [
                        "艾米丽（局促，低声）：hello(你好)",
                        "△ 场景说明（不要删除）",
                        "伊森: 原文对白 (里面有原始括号)(translation)",
                    ]
                ),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        response = await client.post("/api/cleaned-scripts", files=files, data={"title": "修订剧本"})

        assert response.status_code == 201
        payload = response.json()
        assert payload["title"] == "修订剧本"
        assert payload["line_count"] == 3
        assert payload["stripped_count"] == 2
        job_id = payload["id"]

        history_response = await client.get("/api/cleaned-scripts")
        assert history_response.status_code == 200
        assert history_response.json()[0]["id"] == job_id

        detail_response = await client.get(f"/api/cleaned-scripts/{job_id}")
        assert detail_response.status_code == 200
        assert detail_response.json()["cleaned_preview"] == [
            "艾米丽（局促，低声）：hello",
            "△ 场景说明（不要删除）",
            "伊森: 原文对白 (里面有原始括号)",
        ]

        download_response = await client.get(f"/api/cleaned-scripts/{job_id}/download")
        assert download_response.status_code == 200
        assert download_response.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        downloaded = Document(BytesIO(download_response.content))
        assert [paragraph.text for paragraph in downloaded.paragraphs] == [
            "艾米丽（局促，低声）：hello",
            "△ 场景说明（不要删除）",
            "伊森: 原文对白 (里面有原始括号)",
        ]

        scripts_response = await client.get("/api/scripts")
        assert scripts_response.status_code == 200
        assert scripts_response.json()[0]["title"] == "既有剧本"

    import sqlite3

    with sqlite3.connect(Path(tmp_path) / "app.db") as connection:
        assert connection.execute("select count(*) from scripts").fetchone()[0] == 1
        assert connection.execute("select count(*) from cleaned_script_jobs").fetchone()[0] == 1
        assert connection.execute("select count(*) from schema_migrations").fetchone()[0] == 1

    app.state.db._initialized = False
    async with AsyncClient(
        transport=ASGITransport(app=app),
        base_url="http://testserver",
    ) as client:
        assert (await client.get("/api/cleaned-scripts")).status_code == 200

    with sqlite3.connect(Path(tmp_path) / "app.db") as connection:
        assert connection.execute("select count(*) from scripts").fetchone()[0] == 1
        assert connection.execute("select count(*) from cleaned_script_jobs").fetchone()[0] == 1
        assert connection.execute("select count(*) from schema_migrations").fetchone()[0] == 1
