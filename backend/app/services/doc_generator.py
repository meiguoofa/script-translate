from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DEFAULT_ASCII_FONT = "Calibri"
DEFAULT_EAST_ASIA_FONT = "Microsoft YaHei"
DEFAULT_COMPLEX_SCRIPT_FONT = "Arial"
DEFAULT_FONT_SIZE_PT = 11


def _set_rfonts(element, *, ascii_font: str, east_asia_font: str, complex_script_font: str) -> None:
    r_fonts = element.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        element.append(r_fonts)

    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    r_fonts.set(qn("w:cs"), complex_script_font)


def _apply_document_defaults(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = DEFAULT_ASCII_FONT
    normal_style.font.size = Pt(DEFAULT_FONT_SIZE_PT)
    _set_rfonts(
        normal_style.element.get_or_add_rPr(),
        ascii_font=DEFAULT_ASCII_FONT,
        east_asia_font=DEFAULT_EAST_ASIA_FONT,
        complex_script_font=DEFAULT_COMPLEX_SCRIPT_FONT,
    )


def _add_text_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = DEFAULT_ASCII_FONT
    run.font.size = Pt(DEFAULT_FONT_SIZE_PT)
    _set_rfonts(
        run._element.get_or_add_rPr(),
        ascii_font=DEFAULT_ASCII_FONT,
        east_asia_font=DEFAULT_EAST_ASIA_FONT,
        complex_script_font=DEFAULT_COMPLEX_SCRIPT_FONT,
    )


def generate_docx(rendered_lines: list[str], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _apply_document_defaults(document)
    for line in rendered_lines:
        _add_text_paragraph(document, line)
    document.save(output_path)
